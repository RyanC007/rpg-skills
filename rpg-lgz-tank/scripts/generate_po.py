#!/usr/bin/env python3
"""
Tank — Logoclothz Purchase Order (PO) Generator
=================================================
Takes a BigCommerce Order ID and generates a formatted Purchase Order
in .docx format based on a template.

Usage:
    python3 generate_po.py --order-id 4203897 --template-path ../templates/po_template.docx --output-path /tmp/PO_4203897.docx
"""

import requests
import json
import argparse
from datetime import datetime
from docx import Document

# ─────────────────────────────────────────────
# CREDENTIALS & CONFIG
# ─────────────────────────────────────────────
BC_STORE = "dw57ootmu7"
BC_TOKEN = "o1dhpea9cz5zuvqjcr8cc6imoufzcxr"
HEADERS = {
    "X-Auth-Token": BC_TOKEN,
    "Accept": "application/json",
    "Content-Type": "application/json"
}
BASE_URL = f"https://api.bigcommerce.com/stores/{BC_STORE}/v2"

# ─────────────────────────────────────────────
# API FETCH FUNCTIONS
# ─────────────────────────────────────────────
def fetch_order_data(order_id: int) -> dict:
    """Fetch the main order object."""
    resp = requests.get(f"{BASE_URL}/orders/{order_id}", headers=HEADERS)
    if resp.status_code == 200:
        return resp.json()
    print(f"[ERROR] Could not fetch order {order_id}. Status: {resp.status_code}")
    return {}

def fetch_order_products(order_id: int) -> list:
    """Fetch product line items for the order."""
    resp = requests.get(f"{BASE_URL}/orders/{order_id}/products", headers=HEADERS)
    return resp.json() if resp.status_code == 200 else []

def fetch_shipping_address(order_id: int) -> dict:
    """Fetch the primary shipping address for the order."""
    resp = requests.get(f"{BASE_URL}/orders/{order_id}/shipping_addresses", headers=HEADERS)
    if resp.status_code == 200 and resp.json():
        return resp.json()[0]
    return {}

# ─────────────────────────────────────────────
# DOCX HELPER FUNCTIONS
# ─────────────────────────────────────────────
def replace_text_in_doc(doc, replacements: dict):
    """Replace all placeholder keys with their values throughout the document."""
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                # This is a simple text replacement. For more complex formatting,
                # a different approach with runs would be needed.
                inline = p.runs
                # Replace strings and retain formatting
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, str(value))
                        inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

# ─────────────────────────────────────────────
# MAIN EXECUTION
# ─────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="Generate a .docx Purchase Order from a BigCommerce order.")
    parser.add_argument("--order-id", required=True, type=int, help="The BigCommerce Order ID.")
    parser.add_argument("--template-path", required=True, help="Path to the .docx template file.")
    parser.add_argument("--output-path", required=True, help="Path to save the generated .docx file.")
    args = parser.parse_args()

    print(f"[Tank] Generating PO for Order #{args.order_id}...")

    # 1. Fetch all required data
    order = fetch_order_data(args.order_id)
    products = fetch_order_products(args.order_id)
    shipping = fetch_shipping_address(args.order_id)

    if not order:
        return

    billing = order.get("billing_address", {})

    # 2. Prepare data replacements
    replacements = {
        "{{ORDER_ID}}": str(order.get("id", ""))
        ,"{{PO_DATE}}": datetime.now().strftime("%m/%d/%Y"),
        "{{CUSTOMER_NAME}}": f"{billing.get("first_name", "")} {billing.get("last_name", "")}".strip(),
        "{{CUSTOMER_EMAIL}}": billing.get("email", ""),
        "{{CUSTOMER_COMPANY}}": billing.get("company", ""),
        "{{CUSTOMER_PHONE}}": billing.get("phone", ""),
        "{{SHIPPING_ADDRESS}}": f"{shipping.get("street_1", "")} {shipping.get("street_2", "")}\n{shipping.get("city", "")}, {shipping.get("state", "")} {shipping.get("zip", "")}\n{shipping.get("country", "")}".strip(),
        "{{SUBTOTAL}}": f"${float(order.get("subtotal_ex_tax", 0)):.2f}",
        "{{SHIPPING_COST}}": f"${float(order.get("shipping_cost_ex_tax", 0)):.2f}",
        "{{TOTAL_TAX}}": f"${float(order.get("total_tax", 0)):.2f}",
        "{{ORDER_TOTAL}}": f"${float(order.get("total_inc_tax", 0)):.2f}",
        "{{NOTES}}": order.get("customer_message", "N/A")
    }

    # 3. Load template and perform replacements
    try:
        doc = Document(args.template_path)
    except Exception as e:
        print(f"[ERROR] Failed to load template file: {args.template_path}\n{e}")
        return

    replace_text_in_doc(doc, replacements)

    # 4. Populate the product table (assumes the first table is the product table)
    try:
        table = doc.tables[0]
        # Clear placeholder rows (keep header)
        while len(table.rows) > 1:
            table._element.remove(table.rows[-1]._element)

        # Add a new row for each product
        for p in products:
            row_cells = table.add_row().cells
            row_cells[0].text = p.get("sku", "")
            row_cells[1].text = p.get("name", "")
            row_cells[2].text = str(p.get("quantity", ""))
            row_cells[3].text = f"${float(p.get("price_ex_tax", 0)):.2f}"
            row_cells[4].text = f"${float(p.get("total_ex_tax", 0)):.2f}"
    except IndexError:
        print("[WARNING] Could not find a table in the document to populate with products.")
    except Exception as e:
        print(f"[ERROR] Failed to populate product table: {e}")

    # 5. Save the new document
    try:
        doc.save(args.output_path)
        print(f"[Tank] Successfully generated PO: {args.output_path}")
    except Exception as e:
        print(f"[ERROR] Failed to save generated document: {args.output_path}\n{e}")

if __name__ == "__main__":
    main()
